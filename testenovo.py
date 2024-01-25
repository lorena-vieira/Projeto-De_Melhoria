import tkinter as tk
from tkinter import messagebox
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


import utils
import Tela1


# Exemplo de uso
if __name__ == '__main__':
    root = tk.Tk()
    app = Tela1.DataEntryApp(root)
    root.mainloop()

class RelatorioDiagnosticoApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Relatório de Diagnóstico de Pontes Metálicas Ferroviárias")

        # Variáveis para armazenar dados
        self.km_ponte = tk.StringVar()
        self.cidade_ponte = tk.StringVar()
        self.estado_ponte = tk.StringVar()

        # Variáveis para checkboxes
        self.inadequacoes_contraventamento = tk.BooleanVar()
        self.corrosao_media = tk.BooleanVar()
        self.sujeira_vegetacao = tk.BooleanVar()
        self.desgaste_pintura_corrosao_superficial = tk.BooleanVar()

        # Elementos da interface
        self.label_km_ponte = tk.Label(root, text="Km da Ponte:")
        self.entry_km_ponte = tk.Entry(root, textvariable=self.km_ponte)

        self.label_cidade_ponte = tk.Label(root, text="Cidade:")
        self.entry_cidade_ponte = tk.Entry(root, textvariable=self.cidade_ponte)

        self.label_estado_ponte = tk.Label(root, text="Estado:")
        self.entry_estado_ponte = tk.Entry(root, textvariable=self.estado_ponte)

        self.btn_gerar_relatorio = tk.Button(root, text="Gerar Relatório", command=self.gerar_relatorio)

        # Layout da interface
        self.label_km_ponte.grid(row=0, column=0, padx=10, pady=5, sticky=tk.W)
        self.entry_km_ponte.grid(row=0, column=1, padx=10, pady=5)

        self.label_cidade_ponte.grid(row=1, column=0, padx=10, pady=5, sticky=tk.W)
        self.entry_cidade_ponte.grid(row=1, column=1, padx=10, pady=5)

        self.label_estado_ponte.grid(row=2, column=0, padx=10, pady=5, sticky=tk.W)
        self.entry_estado_ponte.grid(row=2, column=1, padx=10, pady=5)

        self.btn_gerar_relatorio.grid(row=3, column=0, columnspan=2, pady=10)

        # Checkboxes para seleção de inadequações
        self.checkbox_inadequacoes = tk.Checkbutton(root, text="Inadequações de Contraventamento", variable=self.inadequacoes_contraventamento)
        self.checkbox_corrosao_media = tk.Checkbutton(root, text="Corrosão média, com perda de seção de aço", variable=self.corrosao_media)
        self.checkbox_sujeira_vegetacao = tk.Checkbutton(root, text="Sujeira e vegetação", variable=self.sujeira_vegetacao)
        self.checkbox_desgaste_pintura_corrosao_superficial = tk.Checkbutton(root, text="Desgaste da pintura e corrosão superficial", variable=self.desgaste_pintura_corrosao_superficial)

        # Adiciona checkboxes ao layout
        self.checkbox_inadequacoes.grid(row=4, column=0, columnspan=2, pady=5, sticky=tk.W)
        self.checkbox_corrosao_media.grid(row=5, column=0, columnspan=2, pady=5, sticky=tk.W)
        self.checkbox_sujeira_vegetacao.grid(row=6, column=0, columnspan=2, pady=5, sticky=tk.W)
        self.checkbox_desgaste_pintura_corrosao_superficial.grid(row=7, column=0, columnspan=2, pady=5, sticky=tk.W)


    def gerar_relatorio(self):
        # Função que será chamada ao clicar no botão "Gerar Relatório"
        km_ponte = self.km_ponte.get()
        cidade_ponte = self.cidade_ponte.get()
        estado_ponte = self.estado_ponte.get()

        # Caminho do documento existente
        existing_document_path = r'C:\Users\Lorena\PycharmProjects\Projeto-de-melhoria\diagnostic_report.docx'

        # Abrir o documento existente
        doc = Document(existing_document_path)

        # Adicionar informações ao documento
        doc.add_heading(f"Relatório de Diagnóstico para a Ponte situada no Km {km_ponte}", level=1)
        doc.add_paragraph(f"Cidade: {cidade_ponte}")
        doc.add_paragraph(f"Estado: {estado_ponte}")


        # Adicionar textos selecionados com base nas checkboxes
        if self.inadequacoes_contraventamento.get():
            texto_inadequacoes = ("Os contraventamentos possuem configurações e características não recomendadas, como indicado a seguir. "
                                  "As diagonais do contraventamento horizontal possuem índice de esbeltez superior ao recomendado para esses elementos. "
                                  "Além disso não existe um sistema efetivo de contraventamento no nível superior das vigas, o que é mais adequado. "
                                  "Ainda, o contraventamento possui ligações que não respeitam a recomendação mínima de 3 conectores. "
                                  "Essas condições não representam risco à segurança da operação ferroviária. Trata-se, portanto, de uma observação em caráter preventivo.")
            doc.add_paragraph(texto_inadequacoes)

        if self.corrosao_media.get():
            texto_corrosao_media = ("É observada corrosão média nas vigas principais, com perdas de seção de aço, principalmente na região inferior das almas, "
                                    "próximo aos aparelhos de apoio. Cabe salientar, que as perdas na região inferior da alma das vigas principais representam risco à segurança estrutural.")
            doc.add_paragraph(texto_corrosao_media)

        if self.sujeira_vegetacao.get():
            texto_sujeira_vegetacao = ("É possível observar vegetação e sujeira depositada sobre a estrutura metálica. Principalmente nas mesas inferiores, "
                                        "os detritos obstruem a drenagem e geram acúmulo de água.")
            doc.add_paragraph(texto_sujeira_vegetacao)

        if self.desgaste_pintura_corrosao_superficial.get():
            texto_desgaste_pintura_corrosao = "A pintura da ponte encontra-se desgastada, e são observadas regiões de corrosão."
            doc.add_paragraph(texto_desgaste_pintura_corrosao)


        # Salvar o documento modificado
        output_path = f"Relatorio_{km_ponte}.docx"
        doc.save(output_path)

        # Exibir mensagem de sucesso
        messagebox.showinfo("Relatório Gerado", f"Relatório para a Ponte {km_ponte} criado com sucesso! Salvo em {output_path}")



# Criar a instância da aplicação
root = tk.Tk()
app = RelatorioDiagnosticoApp(root)

# Iniciar o loop principal da interface gráfica
root.mainloop()




